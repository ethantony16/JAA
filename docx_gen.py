from docx import Document
from docx.shared import RGBColor
import re
from bs4 import BeautifulSoup

def create_resume_docx(content: str, output_path: str):
    doc = Document()
    
    # Split content and thinking process
    parts = content.split("[[THINKING_PROCESS]]")
    main_content = parts[0].strip()
    thinking_content = parts[1].strip() if len(parts) > 1 else ""
    
    # Simple parser for <ins> and <del> tags
    # We will split by tags and add runs accordingly
    # Parse Markdown-ish structure
    # We look for:
    # ### Heading -> Heading 1 (or 2)
    # **Bold** -> Bold run
    # - Bullet -> List Bullet
    # Normal text -> Normal paragraph
    
    lines = main_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('###'):
            # Heading
            text = line.replace('###', '').strip()
            # Still need to parse <ins>/<del> if present in header, though unlikely for structure
            # For simplicity, let's assume headers don't have diff tags for now, or we strip them
            # Check for <ins>/<del> just in case
            p = doc.add_heading('', level=2)
            _add_rich_text(p, text)
            
        elif line.startswith('- '):
            # Bullet point
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            _add_rich_text(p, text)
            
        else:
            # Normal paragraph
            p = doc.add_paragraph()
            _add_rich_text(p, line)

    # Add Thinking Process on new page
    if thinking_content:
        doc.add_page_break()
        doc.add_heading('LLM Thinking Process', level=1)
        doc.add_paragraph(thinking_content)

    doc.save(output_path)

def _add_rich_text(paragraph, text):
    """
    Parses text for <ins>, <del>, and **bold** tags and adds runs to the paragraph.
    This is a mini-parser.
    """
    # Regex to capture all tags
    # <ins>...</ins>
    # <del>...</del>
    # **...**
    
    # We need a regex that splits by ALL of these
    # Note: Regex split captures groups. 
    # Current limitation: Nested tags (e.g. **<ins>text</ins>**) might fail or look properly handled if we process sequentially.
    # Let's keep it simple: Split by one, then process? No, single regex is better.
    
    # Pattern explanation:
    # (<ins>.*?</ins>) matches ins
    # (<del>.*?</del>) matches del
    # (\*\*.*?\*\*) matches bold
    
    pattern = r'(<ins>.*?</ins>|<del>.*?</del>|\*\*.*?\*\*)'
    tokens = re.split(pattern, text)
    
    for token in tokens:
        if not token: 
            continue
            
        if token.startswith("<ins>"):
            run_text = token.replace("<ins>", "").replace("</ins>", "")
            run = paragraph.add_run(run_text)
            run.font.color.rgb = RGBColor(0, 128, 0)
        elif token.startswith("<del>"):
            run_text = token.replace("<del>", "").replace("</del>", "")
            run = paragraph.add_run(run_text)
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.strike = False
        elif token.startswith("**"):
            run_text = token.replace("**", "")
            run = paragraph.add_run(run_text)
            run.bold = True
        else:
            # Plain text
            paragraph.add_run(token)

def create_simple_docx(content: str, output_path: str):
    doc = Document()
    
    # Split content and thinking process
    parts = content.split("[[THINKING_PROCESS]]")
    main_content = parts[0].strip()
    thinking_content = parts[1].strip() if len(parts) > 1 else ""
    
    # Handle markdown-like headings?
    for line in main_content.split('\n'):
        if line.strip().startswith('**') and line.strip().endswith('**'):
             doc.add_heading(line.strip().replace('**', ''), level=2)
        else:
             doc.add_paragraph(line)
             
    # Add Thinking Process on new page
    if thinking_content:
        doc.add_page_break()
        doc.add_heading('LLM Thinking Process', level=1)
        doc.add_paragraph(thinking_content)
        
    doc.save(output_path)
